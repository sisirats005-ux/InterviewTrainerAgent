import os
import json
import pickle
import faiss
import numpy as np
import logging
from langchain_core.documents import Document

# Configure Logging
logging.basicConfig(
    filename='app.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Define paths
KNOWLEDGE_BASE_DIR = "knowledge_base"
VECTOR_DB_DIR = "vector_db"
INDEX_PATH = os.path.join(VECTOR_DB_DIR, "index.faiss")
DOCS_PATH = os.path.join(VECTOR_DB_DIR, "docs.pkl")
MANIFEST_PATH = os.path.join(VECTOR_DB_DIR, "manifest.json")

# Global lazy loaded embedding model
_encoder = None

# Global in-memory FAISS caches to avoid repeated disk reads on active requests
_faiss_index = None
_docs_cache = None
_db_verified = False  # Set to True once verified on startup to bypass folder scans

def get_encoder():
    """
    Lazy-loads and returns the sentence-transformers model.
    """
    global _encoder
    if _encoder is None:
        logging.info("Loading sentence-transformers embedding model (all-MiniLM-L6-v2)...")
        from sentence_transformers import SentenceTransformer
        _encoder = SentenceTransformer("all-MiniLM-L6-v2")
    return _encoder

def extract_text_from_pdf(pdf_path):
    """
    Extracts text from PDF files using PyMuPDF (fitz).
    """
    try:
        import fitz
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        return text
    except Exception as e:
        logging.error(f"Error reading PDF RAG document {pdf_path}: {e}")
        return ""

def extract_text_from_docx(docx_path):
    """
    Extracts text from Word documents using python-docx.
    """
    try:
        import docx
        doc = docx.Document(docx_path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        return "\n\n".join(text)
    except Exception as e:
        logging.error(f"Error reading DOCX RAG document {docx_path}: {e}")
        return ""

def split_text_into_chunks(text, source_name="", chunk_size=700, chunk_overlap=150):
    """
    Splits text into chunks maintaining paragraph context and custom overlap size.
    """
    paragraphs = text.split("\n\n")
    chunks = []
    current_chunk = ""
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        # If adding paragraph is within limits, add it
        if len(current_chunk) + len(para) + 2 <= chunk_size:
            current_chunk += ("\n\n" if current_chunk else "") + para
        else:
            if current_chunk:
                chunks.append(current_chunk)
            
            # Handle large paragraphs
            if len(para) > chunk_size:
                words = para.split()
                temp_chunk = ""
                # Implement sliding window overlap
                for word in words:
                    if len(temp_chunk) + len(word) + 1 <= chunk_size:
                        temp_chunk += (" " if temp_chunk else "") + word
                    else:
                        if temp_chunk:
                            chunks.append(temp_chunk)
                        # Retain trailing overlap words
                        overlap_words = temp_chunk.split()[-25:] if temp_chunk else []
                        temp_chunk = " ".join(overlap_words) + (" " if overlap_words else "") + word
                current_chunk = temp_chunk
            else:
                current_chunk = para
                
    if current_chunk:
        chunks.append(current_chunk)
        
    return [
        Document(page_content=chunk, metadata={"source": source_name})
        for chunk in chunks
    ]

def get_knowledge_base_manifest():
    """
    Scans the knowledge base directory and generates a metadata dict
    containing file sizes and last modified times.
    """
    if not os.path.exists(KNOWLEDGE_BASE_DIR):
        return {}
        
    manifest = {}
    valid_extensions = {".txt", ".pdf", ".docx", ".md"}
    
    for file_name in os.listdir(KNOWLEDGE_BASE_DIR):
        _, ext = os.path.splitext(file_name)
        if ext.lower() in valid_extensions:
            file_path = os.path.join(KNOWLEDGE_BASE_DIR, file_name)
            stat_info = os.stat(file_path)
            manifest[file_name] = {
                "size": stat_info.st_size,
                "mtime": stat_info.st_mtime
            }
    return manifest

def build_vector_db(force_rebuild=False):
    """
    Compares the current files in knowledge_base against manifest.json.
    Automatically rebuilds the index ONLY if differences are detected.
    Reuses in-memory verification flags to avoid redundant folder scans.
    """
    global _db_verified
    
    # If already verified in this session and not a forced request, bypass manifest checking
    if _db_verified and not force_rebuild:
        return
        
    os.makedirs(VECTOR_DB_DIR, exist_ok=True)
    os.makedirs(KNOWLEDGE_BASE_DIR, exist_ok=True)
    
    current_manifest = get_knowledge_base_manifest()
    
    # Check if index exists and manifest matches
    should_rebuild = force_rebuild
    
    if not should_rebuild:
        if (not os.path.exists(INDEX_PATH) or 
            not os.path.exists(DOCS_PATH) or 
            not os.path.exists(MANIFEST_PATH)):
            should_rebuild = True
        else:
            try:
                with open(MANIFEST_PATH, "r", encoding="utf-8") as f:
                    cached_manifest = json.load(f)
                
                # Check for additions, deletions, or structural modifications
                if cached_manifest != current_manifest:
                    logging.info("Knowledge base file alterations detected. Rebuilding FAISS vector index...")
                    print("Knowledge base alterations detected. Rebuilding vector index...")
                    should_rebuild = True
            except Exception as e:
                logging.warning(f"Error reading manifest.json: {e}. Defaulting to rebuild.")
                should_rebuild = True
                
    if not should_rebuild:
        logging.info("FAISS vector database is synchronized with cache. Skipping rebuild.")
        print("Vector database is up to date (cache hit).")
        _db_verified = True
        return
        
    # Rebuild database
    all_documents = []
    for file_name, meta in current_manifest.items():
        file_path = os.path.join(KNOWLEDGE_BASE_DIR, file_name)
        _, ext = os.path.splitext(file_name)
        ext = ext.lower()
        
        content = ""
        if ext in {".txt", ".md"}:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            except Exception as e:
                logging.error(f"Error reading text document {file_name}: {e}")
        elif ext == ".pdf":
            content = extract_text_from_pdf(file_path)
        elif ext == ".docx":
            content = extract_text_from_docx(file_path)
            
        if content.strip():
            doc_chunks = split_text_into_chunks(content, source_name=file_name)
            all_documents.extend(doc_chunks)
            logging.info(f"Loaded '{file_name}' ({ext}): Created {len(doc_chunks)} chunks.")
            
    if not all_documents:
        logging.warning("No documents found in knowledge base to index.")
        return
        
    # Embed and index
    encoder = get_encoder()
    texts = [doc.page_content for doc in all_documents]
    logging.info(f"Generating embeddings for {len(texts)} document chunks...")
    embeddings = encoder.encode(texts, show_progress_bar=False)
    
    embeddings = np.array(embeddings).astype("float32")
    faiss.normalize_L2(embeddings)
    
    dimension = embeddings.shape[1]
    index = faiss.IndexFlatIP(dimension)
    index.add(embeddings)
    
    # Save files
    try:
        faiss.write_index(index, INDEX_PATH)
        with open(DOCS_PATH, "wb") as f:
            pickle.dump(all_documents, f)
        with open(MANIFEST_PATH, "w", encoding="utf-8") as f:
            json.dump(current_manifest, f, indent=2)
            
        # Invalidate in-memory caches to force reload on next retrieve_context
        global _faiss_index, _docs_cache
        _faiss_index = None
        _docs_cache = None
        _db_verified = True
        
        logging.info("FAISS index and manifest file successfully cached.")
        print("Vector database successfully rebuilt and saved.")
    except Exception as e:
        logging.error(f"Error saving FAISS index: {e}")
        print(f"Error saving FAISS index: {e}")

def retrieve_context(query, k=3):
    """
    Retrieves the top-k relevant documents from the cached index.
    Uses global in-memory caches to avoid repeated disk reads.
    """
    global _faiss_index, _docs_cache
    
    # Ensure vector DB is initialized (will skip checks if already verified)
    build_vector_db()
    
    if not os.path.exists(INDEX_PATH) or not os.path.exists(DOCS_PATH):
        return []
        
    try:
        # Load from disk ONLY if not cached in memory
        if _faiss_index is None:
            logging.info("RAG Cache Miss: Loading FAISS index from disk...")
            _faiss_index = faiss.read_index(INDEX_PATH)
        if _docs_cache is None:
            logging.info("RAG Cache Miss: Loading documents from disk...")
            with open(DOCS_PATH, "rb") as f:
                _docs_cache = pickle.load(f)
                
        encoder = get_encoder()
        query_vector = encoder.encode([query]).astype("float32")
        faiss.normalize_L2(query_vector)
        
        distances, indices = _faiss_index.search(query_vector, k)
        
        retrieved = []
        for idx in indices[0]:
            if idx != -1 and idx < len(_docs_cache):
                retrieved.append(_docs_cache[idx])
        return retrieved
    except Exception as e:
        logging.error(f"Error during context retrieval: {e}")
        return []

if __name__ == "__main__":
    print("Testing upgraded RAG document loader...")
    # Force rebuild to generate new index and manifest
    build_vector_db(force_rebuild=True)
    
    # Add a mock file to test auto-rebuild
    mock_file = os.path.join(KNOWLEDGE_BASE_DIR, "mock_test.txt")
    with open(mock_file, "w", encoding="utf-8") as f:
        f.write("This is a temporary RAG document to test index change triggers.")
        
    print("\nTriggering RAG checker to test automatic rebuild...")
    build_vector_db() # Should detect mock_test.txt and rebuild
    
    # Clean up mock file
    if os.path.exists(mock_file):
        os.remove(mock_file)
        
    print("\nTriggering RAG checker again after removing mock...")
    build_vector_db() # Should rebuild again since a file was deleted
